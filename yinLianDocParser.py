#! /usr/bin/env python3
# -*- coding:utf-8 -*-
import os
import re
import json
import time
import argparse
import encodings.idna
from threading import Thread, Lock
from queue import Queue
from urllib.parse import quote
from urllib import request
import pandas
import docx
from transCoordinateSystem import gcj02_to_wgs84, gcj02_to_bd09
from area_code import area_code
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm, Inches


parent = {
    "id": "-1", "name": "全国", "childs":[]
}


# 通过堆栈来进行模拟上下级
# 捕捉到缩进进行清栈
stack = []

def printStack():
    tmp = ""
    for i in range(len(stack)):
       tmp = tmp + stack[i]["name"] + "=>"
    print(tmp)

def empytStack(count):
    print("clean count ", count, "stackSize", len(stack))
    printStack()

    count2 = 0
    if(len(stack) > 1):
        for i in range(len(stack)):
            global parent
            if(len(stack) == 1):
                break

            tmp = stack.pop()
            parent =  stack[len(stack) - 1]
            parent["childs"].append(tmp)

            count2 += 1

            if(count == 1):
                printStack()
                break


def dump(obj):
  for attr in dir(obj):
    print("obj.%s = %r" % (attr, getattr(obj, attr)))

count = 1
level = 0
if __name__ == "__main__":
    stack.append(parent)

    fn= r'C:/Users/Miste/Desktop/asd.docx'
    doc=docx.Document(fn)
    # for paragraph in doc.paragraphs:
    # 按表格读取全部数据
    for table in doc.tables:
        for row in table.rows:
            # print("333", row.cells[0].text, 123, row.cells[0].paragraphs[0].paragraph_format.first_line_indent)

            indent = row.cells[0].paragraphs[0].paragraph_format.first_line_indent

            if(row.cells[0].text.find("代码表") != -1 or len(row.cells) < 2):
                empytStack(100)
                level = 0

                # if(len(stack) > 1):
                #     empytStack()

                #     data.append(stack.pop())

                #     parent = {"id": "-1", "name": "-1", "childs":[]}
                continue
            # if(row.cells[0].paragraphs[0].alignment == None):
            if(indent == None):
                if(level > 1):
                    empytStack(1)
                level += 1

                tmp = {
                    "id" :row.cells[1].text,
                    "name" :row.cells[0].text,
                    "pid" : parent["id"],
                    "pname": parent["name"],
                    "childs": []
                }
                parent = tmp

                stack.append(parent)

            else:
                tmp = {
                    "id" :row.cells[1].text,
                    "name" :row.cells[0].text,
                    "childs": [],
                    "pid" : parent["id"],
                    "pname": parent["name"]
                }
                parent["childs"].append(tmp)

            printStack()

            # 随机打印, 方便了解进度
            count += 1
            if(count > 100):
                count = 1
                print("==============================>", stack)
        empytStack(100)
        print("==============================>", stack)
